# -- coding: utf-8 --
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt


def run(fileName, dico):
    docIn = Document(fileName)
    template = Document('./../template.docx')
    template.save('./../new.docx')
    docOut = Document('./../new.docx')

    delete = False

    for para in docOut.paragraphs:
        if 'TITRE1' in para.text:
            delete = True
        if delete:
            remove_paragraph(para)
    delete = False
    for table in docOut.tables:
        if 'Table' in table.cell(0, 0).text:
            delete = True
        if delete:
            remove_table(table)


    # for para in template.paragraphs:
    #     if 'TITRE1' in para.text:
    #         titre1_style = docOut.styles.add_style('titre1', WD_STYLE_TYPE.PARAGRAPH)
    #         titre1_style = changeStyle(docOut.styles['Heading 1'], para.style)

    #     elif 'TITRE2' in para.text:
    #         titre2_style = para.style
    #         # titre2_style = changeStyle(docOut.styles['Heading 2'], para.style)

    #     elif 'TEST' in para.text:
    #         test_style = para.style
    #         test_style = changeStyle(docOut.styles['Heading 2'], para.style)

    # table = iter(template.tables)
    # label_style = table.next().style
    # changeStyle(docOut.styles[label_style.name], label_style)

    # statut_style = table.next().style
    # docOut.styles.add_style(statut_style.name, WD_STYLE_TYPE.TABLE)
    # changeStyle(docOut.styles[statut_style.name], statut_style)

    # testTable_style = table.next().style
    # changeStyle(docOut.styles[testTable_style.name], testTable_style)

    for elem in dico:
        docOut.add_paragraph(elem, style='Heading 2')
        table = ''
        for sousElem in dico[elem]:
            if(dico[elem][sousElem] and not
               isinstance(dico[elem][sousElem], dict)):
                docOut.add_paragraph('\t' + sousElem)
                docOut.add_paragraph('\t' + dico[elem][sousElem])
            if isinstance(dico[elem][sousElem], dict):
                if not table:
                    table = docOut.add_table(rows=1, cols=3)
                    table.columns[0].width = 10000
                    table.columns[1].width = 3828800
                    table.columns[2].width = 3828800
                    table.autofit = False
                    hdr_cells = table.rows[0].cells
                    hdr_cells[1].text = 'Action'
                    hdr_cells[2].text = 'Resultat attendu'
                    row_cells = table.add_row().cells
                    i = 1
                    row_cells[0].text = str(i)
                else:
                    row_cells = table.add_row().cells
                    i += 1
                    row_cells[0].text = str(i)

                for index, etape in enumerate(dico[elem][sousElem]):
                    if index == 2:
                        break
                    row_cells[index + 1].text = dico[elem][sousElem][etape]
        table.columns[0].width = 2

    docOut.save('./../wordOut.generated.docx')


def changeStyle(new, old):
    newFont = new.font
    oldFont = old.font
    newFont.size = oldFont.size
    newFont.name = oldFont.name
    newFont.color.rgb = oldFont.color.rgb
    newFont.bold = oldFont.bold
    newFont.italic = oldFont.italic
    return new


def remove_paragraph(para):
    p = para._element
    p.getparent().remove(p)
    p._p = p._element = None


def remove_table(table):
    tbl = table._tbl
    for row in table.rows:
        tr = row._tr
        tbl.remove(tr)
