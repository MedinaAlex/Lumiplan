# coding: utf8

from docx import Document
import json
from collections import OrderedDict
import sys

import testIHM as ihm
import createDoc as crt
import myTest as test


def run():
    reload(sys)
    sys.setdefaultencoding("utf-8")
    f = open('./../paraT.txt', 'w')
    t = open('./../tableT.txt', 'w')
    a = open('./../arbre.txt', 'a')
    dd = open('./../dico.txt', 'w')

    fileName = './../word/forTest.docx'
    document = Document(fileName)  # Document word

    para = document.paragraphs  # Liste des paragraphes du document
    tables = document.tables  # Liste des tables du document
    tt = iter(tables)  # itérateur des tables
    d2 = {}  # Liste de Dictionnaire ordonné contenant le contenu
    d = d2['Fiches'] = []
    dico = OrderedDict()
    pre, tag = [''] * 2
    tagList = OrderedDict()
    element = ('Modifié', 'ID', 'Nature', 'Type', 'Statut',
               'Importance', 'Jalons',
               'ID:', 'Type:', 'Importance:', 'Jalons:')
    replacements = {'é': 'e', 'è': 'e', ':': ''}

    for elem in tables:
        for row in range(len(elem.rows)):
            for col in range(len(elem.columns)):
                t.write(elem.cell(row, col).text)
    t.close()

    for elem in para:
        # print('-' * 60)
        # print(elem.text.encode('utf-8'))

        # if 'Exigence' in elem.text:  # On va écrire les prérequis avant
        #     cell = tt.next()

        #     d[tag].update({cell.row_cells(0)[0].text:
        #                    cell.row_cells(0)[1].text})
        #     f.write(cell.row_cells(0)[0].text + '\n')
        #     f.write('\t' + cell.row_cells(0)[1].text)

        # f.write(elem.text + '\n')

        # Créé est le premier élément d'une fiche de test
        if 'Créé' in elem.text:
            # tag = ' '.join(pre.split()[:2])
            if dico:
                d.append(dico)

            dico = OrderedDict()
            tagList[tag] = OrderedDict()

            tmp = elem.text.split(':')
            dico['Titre'] = pre
            dico['Cree'] = ':'.join(tmp[1:])
            dico['Etapes'] = []

        if 'Exigences' in elem.text.split():
            i = 1
            tmp = ''
            while 'Etape' not in para[para.index(elem) + i].text:
                tmp += '\t\t' + para[para.index(elem) + i].text + '\n'
                i += 1

            dico['Exigences'] = tmp

        if 'Description' in elem.text:
            i = 1
            while not para[para.index(elem) + i].text.split:
                i += 1

            dico['Description'] = para[para.index(elem) + i].text

            cell = tt.next()
            dico['Pre-requis'] = cell.row_cells(0)[0].text

        if 'Etape' in elem.text:
            etape = OrderedDict()
            num = ' '.join(elem.text.split()[1])
            etape['Numero'] = num

            for i in range(3):
                cell = tt.next()
                key = cell.row_cells(0)[0].text

                for src, dest in replacements.iteritems():
                    key = key.split()[0].replace(src, dest)

                etape[key] = cell.row_cells(0)[1].text

                if cell.row_cells(1) and i == 2:
                    key = cell.row_cells(1)[0].text

                    for src, dest in replacements.iteritems():
                        key = key.split()[0].replace(src, dest)

                    etape[key] = cell.row_cells(1)[1].text

                f.write(cell.row_cells(0)[0].text + '\n')
                f.write('\t' + cell.row_cells(0)[1].text + '\n')
            dico['Etapes'].append(etape)

        if any(str in element for str in(elem.text.split())):
            key = elem.text.split()[0]
            for src, dest in replacements.iteritems():
                key = key.split()[0].replace(src, dest)

            tmp = elem.text.split(':')
            dico[key] = ':'.join(tmp[1:])

        pre = elem.text
    d.append(dico)

    json.dump(d2, dd)
    dd.close()

    a.close()
    dd.close()
    f.close()
    print('done')
    l = []
    l = ihm.main(d)
    active = []
    for j in l:
        if isinstance(j, tuple):
            active.append(' '.join(j))
        else:
            active.append(j)

    # crt.run(fileName, d)
    for index, fiche in enumerate(d2['Fiches']):
        if fiche['Titre'] not in active:
            del d2['Fiches'][index]
            continue
        for key1, value1 in fiche.iteritems():
            if isinstance(value1, basestring) and 'Titre' not in key1:
                if fiche['Titre'] + '.' + key1 not in active:
                    del d2['Fiches'][index][key1]
            elif isinstance(value1, list):
                i = -1
                for index2, elem in enumerate(value1):
                    i += 1
                    print(index, i, elem['Numero'])
                    if fiche['Titre'] + '.Etape' + elem['Numero'] not in active:
                        print(i)
                        del d2['Fiches'][index][key1][i]
                        i -= 1


    test.run(d2)
    return (d2, tagList, active)

if __name__ == '__main__':
    run()
