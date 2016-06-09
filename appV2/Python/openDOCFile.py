# coding: utf8
"""
Created : 2016-05-12
@author: Alex Medina
Script qui manipule le fichier docx source pour en créer un dictionnaire
de contenu
"""

from docx import Document
from collections import OrderedDict
import sys

import writeOnTemplate as wrt


def run(fileName):
    """Méthode qui prend en paramètre le chemin du document docx source.
    Cette méthode permet de créer le dictionnaire qui représente tout le
    contenu du fichier source.
    """

    # On définit l'encodage du système comme étant de l'utf-8
    reload(sys)
    sys.setdefaultencoding("utf-8")

    document = Document(fileName)  # Document word

    para = document.paragraphs  # Liste des paragraphes du document
    tables = document.tables  # Liste des tables du document
    tt = iter(tables)  # itérateur des tables
    d2 = {}  # Dictionnaire général
    d = d2['Fiches'] = []  # Liste des fiches
    dico = OrderedDict()  # Dictionnaire ordonné
    pre, tag, parent = [''] * 3

    # Liste des éléments à récupérer dans le texte, drop 'Statut'
    element = ('Modifié', 'ID', 'Nature', 'Type',
               'Importance', 'Jalons',
               'ID:', 'Type:', 'Importance:', 'Jalons:')

    # Liste des caractères spéciaux à remplacer
    replacements = {'é': 'e', 'è': 'e', ':': ''}

    # Pour tous les paragraphes du fichier source
    for elem in para:
        # Récupère le dossier parent
        if 'Heading 1' in elem.style.name:
            # on récupère le dernier dossier
            parent = elem.text.split('>')[-1]
            # On supprime les espaces en début de chaine
            while parent.startswith(' '):
                parent = parent[1:]

        # Créé est le premier élément d'une fiche de test
        if 'Créé' in elem.text:
            # Si le dictionnaire n'est pas vide, on va l'ajouter au
            # dictionnaire général, cela permet le passage des fiches.
            if dico:
                d.append(dico)

            # On créer un nouveau dictionnaire car c'est une nouvelle fiche.
            dico = OrderedDict()

            # On split le paragraphe pour récupérer ce qu'il y a après le ':'
            tmp = elem.text.split(':')
            dico['Titre'] = pre
            dico['Parent'] = parent
            dico['Cree'] = ':'.join(tmp[1:])

            # On créer une liste pour les étapes
            dico['Etapes'] = []

        # Lorsqu'il y a Exigences dans le paragraphe.
        if 'Exigences' in elem.text.split():
            i = 1
            tmp = ''
            # On va récupérer tout le texte jusqu'à la partie des étapes.
            while 'Etape' not in para[para.index(elem) + i].text:
                tmp += para[para.index(elem) + i].text + '\n'
                i += 1

            # On l'enregiste dans le dictionnaire.
            dico['Exigences'] = tmp

        # Lorsqu'il y a Description dans le paragraphe.
        if 'Description' in elem.text:
            i = 2

            # On passe le texte sans caractères
            while para[para.index(elem) + i].text:
                i += 1

            # On enregistre le texte qui nous intéresse.
            dico['Description'] = "\n".join([para[para.index(elem) + i].text for i in range(2, i+1)])

            # On sait que les pré-requis sont après la description.
            # On va donc chercher dans les tables les pré-requis
            cell = tt.next()
            dico['Pre-requis'] = cell.row_cells(0)[1].text

        # Lorqu'il y a Etape dans le paragraphe.
        if 'Etape' in elem.text:
            # Dictionnaire ordonné d'une étape
            etape = OrderedDict()

            # On récupère le numéro de l'étape
            etape['Numero'] = ' '.join(elem.text.split()[1])

            # Il y a 3 tableaux qui va nous intéresser.
            for i in range(3):
                # On récupère le tableau suivant.
                cell = tt.next()

                # La clé est le texte est la première valeur du tableau.
                key = cell.row_cells(0)[0].text

                # On va remplacer les caractères spéciaux.
                for src, dest in replacements.iteritems():
                    key = key.split()[0].replace(src, dest)

                # On ajoute à notre dictionnaire la valeur.
                etape[key] = cell.row_cells(0)[1].text

                # Lors de la 3eme itération, on récupère sur la 2eme ligne.
                if cell.row_cells(1) and i == 2:
                    key = cell.row_cells(1)[0].text

                    # On va remplacer les caractères spéciaux.
                    for src, dest in replacements.iteritems():
                        key = key.split()[0].replace(src, dest)

                    # On ajoute à notre dictionnaire la valeur.
                    etape[key] = cell.row_cells(1)[1].text

            # On ajoute à notre liste d'étapes, l'étape courante.
            dico['Etapes'].append(etape)

        # Pour les autres éléments
        if any(str in element for str in(elem.text.split())):

            # On récupère le 1er mot de l'élément
            key = elem.text.split()[0]

            # On va remplacer les caractères spéciaux.
            for src, dest in replacements.iteritems():
                key = key.split()[0].replace(src, dest)

            # On va ajouter à notre dictionnaire, tout ce qu'il y a après le :
            tmp = elem.text.split(':')
            dico[key] = ':'.join(tmp[1:])

        # La ligne précédente est le texte actuel,
        # cela sert uniquement pour les différentes fiches
        pre = elem.text

    # Lors de la fin des itérations, on ajoute la dernière fiche.
    d.append(dico)

    # On retourne le dictionnaire au complet.
    return d2


def activeDico(d2, l, template, name):
    """Méthode qui prend en paramètre, le dictionnaire complet du fichier.
    Les éléménts activés dans l'ihm, le ficheir template et le chemin
    du fichier de destination.
    """

    # Liste des éléments activés
    active = []
    for j in l:
        if isinstance(j, tuple):
            try:
                active.append(' '.join(j))
            except TypeError:
                t = []
                for elem in j:
                    if isinstance(elem, tuple):
                        t.append('"' + ' '.join(elem) + '"')
                    else:
                        t.append(elem)
                active.append(' '.join(t))
        else:
            active.append(j)

    # On va écrire dans le dictionnaire, les éléments de général et other
    # qui sont activés
    d2['general'] = []
    d2['Other'] = []
    for act in active:
        # if('general' in act.split('.') and ('Etape' and 'Statut') not in act.split('.')):
        if('general' in act.split('.') and
           'Etape' not in act.split('.') and
           'Statut' not in act.split('.')):

            d2['general'].append(act.split('.')[1])

        if 'other' in act.split('.'):
            d2['Other'].append(act.split('.')[1])

    # Pour chaque fiches, on va supprimer les éléments indésirables
    # La liste est inversée pour éviter les problèmes de supression
    for fiche in reversed(d2['Fiches']):
        # On récupère l'index de la fiche courante.
        index = d2['Fiches'].index(fiche)

        if fiche['Parent']:
            if fiche['Parent'] + '.' + fiche['Titre'] not in active:
                d2['Fiches'].pop(index)
        else:
            # Si la fiche n'est pas un élément souhaité, on la suprimme
            if fiche['Titre'] not in active:
                d2['Fiches'].pop(index)

            # On passe à la fiche suivante
            continue

        if 'general.Etape' in active:
            d2['is_Etape'] = True
        else:
            d2['is_Etape'] = False

        if 'general.Statut' in active:
            d2['is_Statut'] = True
        else:
            d2['is_Statut'] = False

    # On ajoute dans le dictionnaire, une liste de 1 au nombre de fiches.
    d2['num'] = [i + 1 for i in range(len(d2['Fiches']))]

    # On appelle la méthode qui va écrire dans le fichier template.
    wrt.run(d2, template, name)
    return (d2)

if __name__ == '__main__':
    run()
